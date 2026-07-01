from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"E:\project5\bug-management-system")
TEMPLATE = Path(r"D:\liulangqixiazhaiyingyong\《鲲鹏云计算技术》项目书模板.docx")
OUT = ROOT / "docs" / "缺陷管理系统项目书-202402210242赵彦文.docx"


def set_run_font(run, size=None, bold=None, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_text(paragraph, text, size=11, bold=False, align=None, color=None, font="宋体"):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font=font)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    return paragraph


def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    set_run_font(run, size=size, bold=True, color=(31, 78, 121))
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25


def add_para_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style:
        p.style = style
    if text:
        set_para_text(p, text)
    return p


def insert_after(paragraph, items):
    anchor = paragraph
    insert_after._num = 0
    for kind, payload in items:
        if kind == "p":
            anchor = add_para_after(anchor, payload)
        elif kind == "h2":
            anchor = add_para_after(anchor)
            set_heading(anchor, payload, 2)
        elif kind == "h3":
            anchor = add_para_after(anchor)
            set_heading(anchor, payload, 3)
        elif kind == "bullet":
            anchor = add_para_after(anchor, "·  " + payload)
            anchor.paragraph_format.left_indent = Inches(0.25)
            anchor.paragraph_format.first_line_indent = Inches(-0.12)
            for run in anchor.runs:
                set_run_font(run, size=11)
        elif kind == "num":
            if not hasattr(insert_after, "_num"):
                insert_after._num = 0
            insert_after._num += 1
            anchor = add_para_after(anchor, f"{insert_after._num}. {payload}")
            anchor.paragraph_format.left_indent = Inches(0.25)
            anchor.paragraph_format.first_line_indent = Inches(-0.12)
            for run in anchor.runs:
                set_run_font(run, size=11)
        elif kind == "code":
            anchor = add_para_after(anchor, payload)
            for run in anchor.runs:
                set_run_font(run, size=9, font="Consolas")
            anchor.paragraph_format.left_indent = Inches(0.25)
        elif kind == "caption":
            anchor = add_para_after(anchor, payload)
            for run in anchor.runs:
                set_run_font(run, size=10, bold=True, color=(89, 89, 89))
            anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return anchor


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    p.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def set_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value, bold=(r == 0), fill=("E8EEF5" if r == 0 else None))


def add_table_after(paragraph, rows, widths=None):
    anchor = add_para_after(paragraph)
    table = anchor._parent.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.4))
    anchor._p.addnext(table._tbl)
    table.style = "Table Grid"
    set_table(table, rows)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def clear_placeholder(paragraph):
    if paragraph.text.strip().startswith("（") or paragraph.text.strip().startswith("内容参考") or paragraph.text.strip().startswith("（内容"):
        set_para_text(paragraph, "")


def main():
    doc = Document(TEMPLATE)

    # Page setup: A4 with moderate margins, matching common Chinese university reports.
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == "项目书":
            set_para_text(p, "项目书", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
        elif txt.startswith("课程名称"):
            set_para_text(p, "课程名称：   鲲鹏云计算技术", size=14)
        elif txt.startswith("学号姓名"):
            set_para_text(p, "学号姓名：   202402210242 赵彦文", size=14)
        elif txt.startswith("专业班级"):
            set_para_text(p, "专业班级：   2024级计算机科学与技术4班", size=14)
        elif txt.startswith("二级学院"):
            set_para_text(p, "二级学院：   电气与信息工程学院", size=14)
        elif txt.startswith("授课老师"):
            set_para_text(p, "授课老师：   段莹莹", size=14)
        elif "2026 年" in txt:
            set_para_text(p, "2026 年 6 月 28 日", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif txt.startswith("2.Docker 容器化部署技术过程"):
            set_para_text(p, "")
        elif txt in {"1项目概述", "2需求分析与设计", "3项目方案设计", "4项目实施过程", "5系统功能测试", "6项目总结", "7附录"}:
            set_heading(p, txt, 1)
        elif txt and txt[0:3] in {"1.1", "1.2", "2.1", "2.2", "2.3", "2.4", "3.1", "3.2", "3.3", "3.4", "3.5", "4.1", "4.2", "4.3", "5.1", "5.2", "5.3"}:
            set_heading(p, txt, 2)

    # Fill the two existing template tables.
    if doc.tables:
        set_table(doc.tables[0], [
            ["功能点", "需求描述", "验收标准"],
            ["用户登录与权限", "提供管理员、测试人员、开发人员三类账号，登录后按角色展示菜单和操作按钮。", "正确账号可登录；无权限功能不显示或不可操作。"],
            ["Bug 缺陷管理", "支持新增、编辑、删除、列表查询、详情查看、标题搜索、状态和严重程度筛选。", "Bug 数据能保存到 MySQL，并可在列表和详情页正确展示。"],
            ["状态流转与记录", "支持已提交、已修复、已关闭三个状态，并记录新增、编辑、转交、关闭等操作。", "状态变更后页面和数据库同步更新，操作日志可追溯。"],
        ])
    if len(doc.tables) > 1:
        set_table(doc.tables[1], [
            ["测试模块", "测试操作", "预期结果", "实际结果", "是否通过"],
            ["用户登录", "输入正确账号密码", "成功进入控制台首页", "admin、tester、developer 均可登录", "通过"],
            ["提交Bug", "填写完整缺陷信息并保存", "数据存入数据库，列表展示", "新增后可在列表和详情页查看", "通过"],
            ["状态流转", "开发人员改为已修复，测试人员改为已关闭", "状态按流程更新并记录日志", "状态和操作记录均正常", "通过"],
            ["Docker访问", "执行 docker compose up -d --build 后访问 8080", "浏览器打开完整系统", "配置已完成，待 Docker Desktop 启动后最终验证", "待环境验证"],
        ])

    # Replace template placeholder paragraphs with full report content.
    content_map = {
        "1.1项目背景": [
            ("p", "随着软件系统功能越来越复杂，测试和开发过程中会不断产生缺陷记录。如果缺陷只依靠聊天工具或表格记录，容易出现处理人不明确、状态不可追踪、历史记录丢失等问题。因此，本项目设计并实现一个基于 Web 的缺陷管理系统，用于完成 Bug 的提交、处理、修复、关闭和统计展示。"),
            ("p", "本项目同时结合《鲲鹏云计算技术》课程要求，将系统部署目标设置为 FusionCompute 虚拟化平台中的鲲鹏 ARM 架构 Linux 虚拟机，并通过 Docker 容器完成项目打包和发布。项目既体现 Web 系统开发能力，也体现虚拟化、容器化和云计算基础环境部署能力。"),
        ],
        "1.2整体实现方案简述": [
            ("p", "整体方案采用“FusionCompute 虚拟机环境 + 前后端分离 Bug 管理系统 + Docker Compose 容器化部署”的方式实现。系统前端使用 Vue 3、Vite 和 Element Plus 开发，提供登录页、控制台、Bug 列表、新增 Bug、详情页、统计分析和用户管理页面；后端使用 Spring Boot 3 开发 REST API，负责登录认证、权限控制、Bug 业务处理和数据库读写；数据库使用 MySQL 8 保存用户、Bug 和操作记录。"),
            ("p", "部署时使用三个容器：mysql、backend、frontend。其中 frontend 容器通过 Nginx 提供静态页面，并将 /api 请求反向代理到 backend 容器；backend 容器通过 JDBC 连接 mysql 容器；mysql 容器在首次启动时执行 init.sql 初始化数据。最终用户通过浏览器访问 http://服务器IP:8080 即可使用系统。"),
        ],
        "2.1功能性需求": [
            ("p", "本系统的功能性需求主要围绕用户、权限、Bug 管理、状态流转、操作记录和统计展示展开。"),
            ("bullet", "用户模块：支持账号密码登录，系统预置 admin、tester、developer 三个用户。"),
            ("bullet", "权限模块：管理员、测试人员、开发人员登录后可见菜单和操作按钮不同。"),
            ("bullet", "Bug 管理模块：支持新增、编辑、删除、查看列表、查看详情、按标题搜索和按状态/严重程度筛选。"),
            ("bullet", "状态流转模块：支持已提交 → 已修复 → 已关闭的缺陷处理流程。"),
            ("bullet", "转交模块：支持将 Bug 从一个处理人转交给另一个处理人，并保存转交备注。"),
            ("bullet", "日志模块：新增、编辑、修改状态、转交和关闭操作均写入操作记录。"),
            ("bullet", "统计模块：展示 Bug 总数、今日新增、未处理、已修复、已关闭和严重 Bug 数量。"),
        ],
        "2.2非功能性需求": [
            ("bullet", "兼容性：系统设计适配 FusionCompute 虚拟机环境，可部署在 openEuler 或鲲鹏适配 Linux 系统上。"),
            ("bullet", "可部署性：项目通过 Dockerfile 和 docker-compose.yml 进行容器化打包，支持一条命令启动。"),
            ("bullet", "易用性：前端页面采用蓝白配色和后台管理布局，表格、按钮和表单清晰，便于课堂演示和项目书截图。"),
            ("bullet", "安全性：接口请求需要携带 token，不同角色在后端进行权限判断，避免越权操作。"),
            ("bullet", "可维护性：前后端目录分离，后端按 controller、service、repository 分层，前端按 api、router、views、layouts 组织。"),
        ],
        "2.3项目约束条件": [
            ("bullet", "硬件环境：课程要求面向鲲鹏 ARM 架构服务器，虚拟化平台使用 FusionCompute。"),
            ("bullet", "系统环境：虚拟机操作系统为 openEuler 或鲲鹏适配 Linux。"),
            ("bullet", "部署技术：最终项目必须通过 Docker 容器方式发布，并使用 docker-compose 进行编排。"),
            ("bullet", "开发周期：项目开发按照需求分析、数据库设计、后端开发、前端开发、联调测试、Docker 化、文档整理几个阶段完成。"),
            ("bullet", "功能范围：优先保证系统可运行和核心功能完整，不做过度复杂的企业级流程设计。"),
        ],
        "2.4项目验收标准": [
            ("p", "项目验收以“能运行、功能完整、容器部署成功、文档齐全”为主要标准。具体包括：系统能够登录；能够新增、查看、搜索和编辑 Bug；能够进行状态流转和转交处理人；能够查看操作记录和统计图表；能够通过 docker compose up -d --build 启动；项目书、部署说明和截图清单完整。"),
        ],
        "3.1总体架构设计": [
            ("p", "系统采用三层架构设计：底层为 FusionCompute 虚拟化层，提供鲲鹏 ARM 架构虚拟机；中间为 Docker 容器部署层，运行 mysql、backend、frontend 三个容器；上层为应用层，向用户提供浏览器访问的 Bug 管理系统。"),
            ("code", "浏览器用户\n   │ 访问 http://服务器IP:8080\n   ▼\nfrontend 容器：Nginx + Vue 静态页面\n   │ /api 反向代理\n   ▼\nbackend 容器：Spring Boot REST API\n   │ JDBC 连接\n   ▼\nmysql 容器：MySQL 8 数据库"),
            ("caption", "图3-1 系统总体架构图"),
            ("p", "前后端分离后，前端主要负责页面展示和用户交互，后端主要负责业务逻辑、权限判断和数据库操作。这种结构便于开发调试，也适合用 Docker 将不同服务分别打包。"),
        ],
        "3.2技术选型": [
            ("p", "本项目的技术选型以课程要求、开发效率和部署可行性为依据。"),
            ("bullet", "虚拟化平台：FusionCompute，用于创建和管理鲲鹏云虚拟机。"),
            ("bullet", "操作系统：openEuler / 鲲鹏适配 Linux，适合 ARM aarch64 服务器环境。"),
            ("bullet", "后端技术：Spring Boot 3，生态成熟，适合开发 REST 接口。"),
            ("bullet", "前端技术：Vue 3 + Vite + Element Plus，开发效率高，适合管理后台页面。"),
            ("bullet", "数据库：MySQL 8，用于保存用户、缺陷和操作记录。"),
            ("bullet", "容器技术：Docker + Docker Compose，分别部署前端、后端和数据库。"),
            ("bullet", "Web 服务：Nginx，用于部署前端静态资源并代理后端接口。"),
        ],
        "3.3数据库设计": [
            ("p", "数据库名称为 bug_management，核心包括 users、bugs、bug_operation_logs 三张表。users 保存用户账号和角色；bugs 保存缺陷主体信息；bug_operation_logs 保存缺陷操作历史。"),
            ("code", "users 1 ── n bugs（creator_id）\nusers 1 ── n bugs（assignee_id）\nbugs  1 ── n bug_operation_logs\nusers 1 ── n bug_operation_logs（operator_id）"),
            ("caption", "图3-2 数据库 E-R 关系简图"),
            ("p", "users 表字段包括 id、username、password、nickname、role、created_at、updated_at。bugs 表字段包括 id、title、module_name、severity、description、reproduce_steps、screenshot_url、status、assignee_id、creator_id、created_at、updated_at。bug_operation_logs 表字段包括 id、bug_id、operator_id、operation_type、old_status、new_status、old_assignee_id、new_assignee_id、remark、created_at。"),
            ("p", "系统初始化时预置三个账号：admin/123456、tester/123456、developer/123456，并插入若干示例 Bug 数据，便于系统启动后直接演示。"),
        ],
        "3.4功能模块详细设计": [
            ("h3", "用户登录权限模块"),
            ("p", "用户输入账号密码后，后端校验 users 表中的账号信息。校验成功后返回 token 和用户信息，前端保存 token，并在之后的请求中放入 Authorization 请求头。系统根据用户角色控制菜单显示和按钮权限。"),
            ("h3", "Bug 管理与状态流转模块"),
            ("p", "测试人员或管理员新增 Bug 后，默认状态为“已提交”。开发人员处理后可将状态改为“已修复”。测试人员复测通过后可将状态改为“已关闭”。管理员可以根据实际情况修改所有状态。每次状态变化都会写入操作记录。"),
            ("code", "新增 Bug → 已提交 → 开发人员处理 → 已修复 → 测试人员复测 → 已关闭"),
            ("h3", "查询统计模块"),
            ("p", "Bug 列表支持按标题、状态、严重程度和处理人筛选。统计模块根据 bugs 表中的状态、严重程度和创建时间进行聚合，向前端返回总数、今日新增、未处理、已修复、已关闭和严重 Bug 数量。"),
        ],
        "3.5容器部署方案": [
            ("p", "项目通过 Docker Compose 编排三个服务：mysql、backend、frontend。mysql 使用官方 mysql:8.0 镜像，并挂载 mysql/init.sql 作为初始化脚本；backend 使用多阶段 Dockerfile 构建 Spring Boot jar 包；frontend 使用 Node 镜像构建 Vue 项目，再使用 Nginx 镜像提供静态页面。"),
            ("code", "docker compose up -d --build\n\ndocker compose ps\ndocker compose logs -f\ndocker compose down"),
            ("p", "端口映射方面，frontend 对外暴露 8080，backend 暴露 8081，mysql 暴露 3306。正常使用时用户只需要访问 8080 端口。Nginx 将 /api 请求代理到 backend:8081，保证浏览器访问入口统一。"),
        ],
        "4.1 FusionCompute虚拟化平台部署操作": [
            ("num", "登录 FusionCompute 管理平台，进入虚拟机资源管理页面。"),
            ("num", "创建一台鲲鹏 ARM 架构 Linux 虚拟机，操作系统选择 openEuler 或鲲鹏适配 Linux。"),
            ("num", "配置虚拟机 CPU、内存、磁盘和网络，保证 8080 端口可访问。"),
            ("num", "通过远程终端连接虚拟机，执行 cat /etc/os-release 和 uname -m 检查系统信息。"),
            ("num", "安装 git、curl、vim、firewalld 等基础工具，为 Docker 部署做准备。"),
        ],
        "4.2 Bug管理系统前后端开发实现": [
            ("p", "数据库部分编写 mysql/init.sql，创建 users、bugs、bug_operation_logs 三张表，并初始化测试账号和示例数据。"),
            ("p", "后端部分使用 Spring Boot 3 开发，完成统一返回格式、登录接口、用户接口、Bug 管理接口、状态流转接口、转交接口、操作记录接口和统计接口。后端通过 JDBC Template 操作 MySQL，并在 Service 层完成角色权限判断。"),
            ("p", "前端部分使用 Vue 3、Vite 和 Element Plus 开发，完成登录页、后台主布局、控制台、Bug 列表、新增/编辑页面、详情页、统计页和用户列表页。前端使用 Axios 封装接口请求，并通过路由守卫控制登录状态。"),
            ("p", "联调阶段使用本地 MySQL 和后端服务验证了登录、获取用户、Bug 增删改查、状态修改、转交处理人、操作记录和统计接口，核心功能均能正常响应。"),
        ],
        "4.3 Docker容器化打包和部署": [
            ("num", "后端 Dockerfile 使用 maven:3.9-eclipse-temurin-17 构建 jar 包，再使用 eclipse-temurin:17-jre 运行。"),
            ("num", "前端 Dockerfile 使用 node:lts-alpine 执行 npm ci 和 npm run build，再使用 nginx:alpine 部署 dist 静态资源。"),
            ("num", "docker-compose.yml 编排 mysql、backend、frontend 三个服务，并配置数据库环境变量、端口映射、数据卷和健康检查。"),
            ("num", "启动命令为 docker compose up -d --build，启动后通过 http://localhost:8080 或 http://服务器IP:8080 访问系统。"),
            ("num", "部署完成后通过 docker compose ps 查看容器状态，通过 docker compose logs -f 查看运行日志。"),
        ],
        "5.1测试环境": [
            ("p", "本项目测试环境包括本地开发环境和目标部署环境。开发环境为 Windows 主机，安装 JDK、Maven、Node.js、MySQL 8 和 Docker 工具；目标部署环境为 FusionCompute 虚拟化平台中的 openEuler / ARM aarch64 虚拟机。浏览器使用 Chrome 或 Edge。"),
            ("p", "阶段性测试中，后端在 8081 端口运行，前端开发服务在 5173 端口运行，前端通过 Vite 代理访问后端 API。Docker 部署时统一通过 8080 端口访问前端 Nginx。"),
        ],
        "5.2功能测试用例": [
            ("p", "功能测试主要围绕登录、Bug 提交、列表查询、状态流转、转交处理人、操作记录和统计展示进行。模板中的测试用例表已填写主要测试项，测试结果表明系统核心业务流程能够正常完成。"),
        ],
        "5.3测试结论": [
            ("p", "经过本地联调测试，系统登录、用户信息获取、Bug 新增、编辑、查询、状态流转、转交处理人、操作日志和统计接口均能正常工作。前端页面能够通过代理访问后端接口，页面结构完整，适合项目展示和截图。Docker 配置已完成，等待 Docker Desktop 或服务器 Docker 环境启动后可执行完整容器部署验证。"),
        ],
        "6项目总结": [
            ("p", "本项目完成了一个基于 Vue 3、Spring Boot 3 和 MySQL 的缺陷管理系统，实现了课程要求中的用户登录与权限、Bug 核心管理、状态流转、转交处理、操作记录和统计图表等功能。项目结构清晰，前后端分离，能够满足学校期末作业的基本要求。"),
            ("p", "在 Docker 容器化方面，本项目通过后端 Dockerfile、前端 Dockerfile、Nginx 配置和 docker-compose.yml 将系统拆分为 mysql、backend、frontend 三个服务，实现了从代码到容器运行的完整部署流程。通过本项目，我进一步理解了容器镜像构建、容器编排、端口映射、服务依赖和数据卷初始化等概念。"),
            ("p", "不足之处是当前系统仍属于课程项目级别，密码采用简单明文存储，文件上传功能使用截图地址代替真实上传，后续可以继续完善密码加密、分页查询、真实附件上传和更细粒度的权限控制。"),
        ],
        "7附录": [
            ("p", "附录部分主要放置关键源码说明和项目运行截图。实际提交项目书时，可根据截图清单补充登录页、控制台、Bug 列表、新增 Bug、详情页、操作记录、Docker 容器运行状态和服务器访问结果截图。"),
        ],
        "源码": [
            ("code", "后端入口：backend/src/main/java/com/example/bugmanagement/BugManagementApplication.java\n登录接口：backend/src/main/java/com/example/bugmanagement/controller/AuthController.java\nBug接口：backend/src/main/java/com/example/bugmanagement/controller/BugController.java\n前端入口：frontend/src/main.js\n前端路由：frontend/src/router/index.js\nDocker编排：docker-compose.yml"),
        ],
        "结果图": [
            ("p", "建议粘贴以下运行截图：登录页、控制台首页、Bug 列表页、新增 Bug 页面、Bug 详情页、状态流转弹窗、操作记录时间线、统计分析页、docker compose ps 容器状态、服务器浏览器访问 http://服务器IP:8080 的结果。"),
        ],
    }

    # Remove simple placeholder guide lines and insert content below the matching headings.
    for p in doc.paragraphs:
        clear_placeholder(p)

    for p in list(doc.paragraphs):
        key = p.text.strip()
        if key in content_map:
            insert_after(p, content_map[key])

    # Add a compact technical table after 3.2 heading.
    for p in doc.paragraphs:
        if p.text.strip() == "3.2技术选型":
            add_table_after(p, [
                ["类别", "选型", "原因"],
                ["前端", "Vue 3 + Vite + Element Plus", "适合后台管理系统，开发效率高，组件丰富。"],
                ["后端", "Spring Boot 3", "接口开发方便，结构清晰，适合 Java Web 项目。"],
                ["数据库", "MySQL 8", "课程常用数据库，便于初始化和部署。"],
                ["部署", "Docker Compose", "可一次启动前端、后端和数据库三个服务。"],
            ], widths=[1.1, 2.1, 3.1])
            break

    # Footer.
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        set_para_text(footer, "缺陷管理系统项目书 - 202402210242 赵彦文", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
